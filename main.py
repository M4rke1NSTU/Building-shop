import os
from flask import Flask, render_template, redirect, url_for, request, session, send_file, jsonify, flash
import io
#from docx import Document
from docx.shared import Inches
from captcha.image import ImageCaptcha
import random
import string
from flask import Flask, render_template, redirect, url_for, request, session, make_response
import sqlite3
import jwt
from datetime import datetime, timedelta
from io import BytesIO
from PIL import Image



app = Flask(__name__)
app.secret_key = 'your_secret_key'
UPLOAD_FOLDER = 'static/avatars'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


def generate_random_text():
    # Определяем доступные символы
    charset = string.ascii_letters + string.digits

    # Генерируем случайный текст длиной 6 символов
    result = ''.join(random.choices(charset, k=6))

    return result

random_text = generate_random_text()

captcha = ImageCaptcha(width=800, height=600)



data = captcha.generate(random_text)

captcha.write(random_text, "captcha.png")


def generate_jwt(user_id):
    payload = {
        'user_id': user_id,
        'exp': datetime.utcnow() + timedelta(minutes=15)
    }
    token = jwt.encode(payload, app.secret_key, algorithm='HS256')
    return token

def generate_refresh_token(user_id):
    payload = {
        'user_id': user_id,
        'exp': datetime.utcnow() + timedelta(days=7)
    }
    refresh_token = jwt.encode(payload, app.secret_key, algorithm='HS256')
    return refresh_token

def verify_jwt(token):
    try:
        payload = jwt.decode(token, app.secret_key, algorithms=['HS256'])
        return payload['user_id']
    except jwt.ExpiredSignatureError:
        return None
    except jwt.InvalidTokenError:
        return None

def get_connection(database):
    conn = sqlite3.connect(database, timeout=10)
    return conn

def create_tables(conn):
    c = conn.cursor()
    c.execute("""CREATE TABLE IF NOT EXISTS Users (
                    user_id INTEGER PRIMARY KEY, 
                    username TEXT, 
                    password TEXT, 
                    email TEXT,
                    refresh_token TEXT)""")
    c.execute("""CREATE TABLE IF NOT EXISTS Products (
                    product_id INTEGER PRIMARY KEY, 
                    type TEXT, 
                    brand TEXT, 
                    model TEXT, 
                    price REAL, 
                    imageURL TEXT, 
                    quantity INTEGER, 
                    isExist BOOLEAN)""")
    c.execute("""CREATE TABLE IF NOT EXISTS Cart (
                    cart_id INTEGER PRIMARY KEY, 
                    user_id INTEGER, 
                    product_id INTEGER, 
                    quantity INTEGER, 
                    FOREIGN KEY (user_id) REFERENCES Users(user_id), 
                    FOREIGN KEY (product_id) REFERENCES Products(product_id))""")
    c.execute("""CREATE TABLE IF NOT EXISTS Orders (
                    order_id INTEGER PRIMARY KEY, 
                    user_id INTEGER, 
                    product_id INTEGER, 
                    quantity INTEGER, 
                    name TEXT, 
                    address TEXT, 
                    phone TEXT, 
                    payment_method TEXT, 
                    FOREIGN KEY (user_id) REFERENCES Users(user_id), 
                    FOREIGN KEY (product_id) REFERENCES Products(product_id))""")

    conn.commit()

    # Check if the refresh_token column exists and add it if not
    try:
        c.execute('SELECT refresh_token FROM Users LIMIT 1')
    except sqlite3.OperationalError:
        c.execute('ALTER TABLE Users ADD COLUMN refresh_token TEXT')
        conn.commit()

    conn.close()

create_tables(get_connection('building_materials_shop.db'))



class Product():
    def __init__(self,product_id, type, brand, model, price, imageURL, quantity, isExist):
        self.product_id = product_id
        self.type = type
        self.brand = brand
        self.model = model
        self.price = int(price)
        self.imageURL = imageURL
        self.quantity = int(quantity)
        self.isExist = isExist




# Маршруты
@app.route('/')
def index():
    access_token = request.cookies.get('token')
    user_id = verify_jwt(access_token)
    conn = get_connection('building_materials_shop.db')
    c = conn.cursor()
    c.execute("SELECT * FROM Products")
    products = [Product(row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7]) for row in c.fetchall()]
    if user_id:
        return render_template('index.html', products=products, current_user=user_id)
    else:
        refresh_token = request.cookies.get('refresh_token')
        user_id = verify_jwt(refresh_token)
        if user_id:
            new_access_token = generate_jwt(user_id)
            new_refresh_token = generate_refresh_token(user_id)
            conn = get_connection('building_materials_shop.db')
            c = conn.cursor()
            print(new_refresh_token, new_access_token)
            c.execute('UPDATE Users SET refresh_token = ? WHERE user_id = ?', (new_refresh_token, user_id))
            conn.commit()
            conn.close()
            response = make_response(redirect(url_for('index')))
            response.set_cookie('token', new_access_token, httponly=True)
            response.set_cookie('refresh_token', new_refresh_token, httponly=True)
            return response
        else:
            return render_template('index.html', products=products)


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()
        email = request.form['email']
        password = request.form['password']
        user = c.execute('SELECT * FROM Users WHERE email = ?', (email,)).fetchone()
        if user and user[2] == password:
            token = generate_jwt(user[0])
            refresh_token = generate_refresh_token(user[0])
            c.execute('UPDATE Users SET refresh_token = ? WHERE user_id = ?', (refresh_token, user[0]))
            conn.commit()
            session['user_id'] = user[0]  # Save user_id in session
            response = make_response(redirect(url_for('index')))
            response.set_cookie('token', token, httponly=True)
            response.set_cookie('refresh_token', refresh_token, httponly=True)
            return response
        else:
            return render_template('login.html', error='Invalid email or password')
    return render_template('login.html')



def madeCaptcha():
    captcha_text = generate_random_text()
    captcha = ImageCaptcha(width=200, height=70)
    data = captcha.generate(captcha_text)
    img = Image.open(BytesIO(data.getvalue()))
    img.save('captcha.png')
    return captcha_text

@app.route('/register', methods=['GET', 'POST'])
def register():
    captcha_text = generate_random_text()
    captcha = ImageCaptcha(width=200, height=70)
    data = captcha.generate(captcha_text)
    img = Image.open(BytesIO(data.getvalue()))
    img.save('captcha.png')
    print(captcha_text)
        # Генерируем новую капчу
    if request.method == 'POST':
        captcha_response = request.form['captcha']
        if captcha_response.lower() != captcha_text.lower():
            return render_template('register.html', error='Неверный код с картинки')
        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        c.execute('INSERT INTO Users (username, email, password) VALUES (?, ?, ?)', (username, email, password))
        conn.commit()
        conn.close()
        return redirect(url_for('login'))

    return render_template('register.html', captcha_img='captcha.png')

@app.route('/profile')
def profile():
    access_token = request.cookies.get('token')
    user_id = verify_jwt(access_token)
    if user_id:
        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()
        c.execute('SELECT username, email FROM Users WHERE user_id = ?', (user_id,))
        user_info = c.fetchone()
        count_of_orders = c.execute("SELECT COUNT(*) FROM Orders WHERE user_id = ?",(user_id,)).fetchone()
        conn.close()
        if user_info:
            user = {'username': user_info[0], 'email': user_info[1]}
            return render_template('profile.html', user=user, orders=count_of_orders[0])
        else:
            return redirect(url_for('index'))
    else:
        # Check if the refresh token is still valid
        refresh_token = request.cookies.get('refresh_token')
        user_id = verify_jwt(refresh_token)
        if user_id:
            # Generate new access token and refresh token
            new_access_token = generate_jwt(user_id)
            new_refresh_token = generate_refresh_token(user_id)
            conn = get_connection('building_materials_shop.db')
            c = conn.cursor()
            c.execute('UPDATE Users SET refresh_token = ? WHERE user_id = ?', (new_refresh_token, user_id))
            conn.commit()
            conn.close()
            response = make_response(redirect(url_for('profile')))
            response.set_cookie('token', new_access_token, httponly=True)
            response.set_cookie('refresh_token', new_refresh_token, httponly=True)
            return response
        else:
            return redirect(url_for('login'))


@app.route('/order_history')
def order_history():
    user_id = session['user_id']
    conn = get_connection('building_materials_shop.db')
    c = conn.cursor()
    c.execute('SELECT p.type, p.brand, p.model, p.price, o.quantity, (p.price * o.quantity) as total '
              'FROM Orders o '
              'JOIN Products p ON o.product_id = p.product_id '
              'WHERE o.user_id = ? '
              'ORDER BY o.order_id DESC', (user_id,))
    orders = [{'type': row[0], 'brand': row[1], 'model': row[2], 'price': row[3], 'quantity': row[4], 'total': row[5]} for row in c.fetchall()]
    conn.close()
    return render_template('orders_history.html', orders=orders)



@app.route('/update_profile', methods=['POST'])
def update_profile():
    if 'user_id' in session:
        user_id = session['user_id']
        username = request.form['username']
        email = request.form['email']

        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()
        c.execute("UPDATE Users SET username = ?, email = ? WHERE user_id = ?", (username, email, user_id))
        conn.commit()
        conn.close()

        # Обновляем информацию о пользователе в сессии
        session['username'] = username

        return redirect(url_for('profile'))
    else:
        return redirect(url_for('login'))


@app.route('/add_to_cart', methods=['POST'])
def add_to_cart():
    access_token = request.cookies.get('token')
    user_id = verify_jwt(access_token)
    if user_id:
        user_id = session['user_id']
        product_id = request.form['product_id']
        quantity = int(request.form['quantity'])
        print(f"User ID: {user_id}, Product ID: {product_id}, Quantity: {quantity}")
        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()

        # Проверяем, есть ли уже этот товар в корзине пользователя
        c.execute('SELECT * FROM Cart WHERE user_id = ? AND product_id = ?', (user_id, product_id))
        existing_item = c.fetchone()
        if existing_item:
            # Если товар уже есть, обновляем количество
            new_quantity = existing_item[3] + quantity
            c.execute('UPDATE Cart SET quantity = ? WHERE cart_id = ?', (new_quantity, existing_item[0]))
        else:
            # Если товара нет, добавляем его в корзину
            c.execute('INSERT INTO Cart (user_id, product_id, quantity) VALUES (?, ?, ?)', (user_id, product_id, quantity))

        conn.commit()
        conn.close()
        return redirect(url_for('index'))
    else:
        refresh_token = request.cookies.get('refresh_token')
        user_id = verify_jwt(refresh_token)
        if user_id:
            new_access_token = generate_jwt(user_id)
            new_refresh_token = generate_refresh_token(user_id)
            conn = get_connection('building_materials_shop.db')
            c = conn.cursor()
            c.execute('UPDATE Users SET refresh_token = ? WHERE user_id = ?', (new_refresh_token, user_id))
            conn.commit()
            conn.close()
            response = make_response(redirect(url_for('index')))
            response.set_cookie('token', new_access_token, httponly=True)
            response.set_cookie('refresh_token', new_refresh_token, httponly=True)
            return response
        else:
            return render_template('login.html')

@app.route('/cart')
def cart():
    token = request.cookies.get('token')
    user_id = verify_jwt(token)
    if user_id:
        user_id = session['user_id']
        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()

        # Получаем все товары в корзине пользователя
        c.execute('SELECT p.type, p.brand, p.model, p.price, p.imageURL, c.quantity, (p.price * c.quantity) as total, p.product_id '
                  'FROM Cart c '
                  'JOIN Products p ON c.product_id = p.product_id '
                  'WHERE c.user_id = ?', (user_id,))
        cart_items = [{'type': row[0], 'brand': row[1], 'model': row[2], 'price': row[3], 'imageURL': row[4], 'quantity': row[5], 'total': row[6], 'id': row[7]} for row in c.fetchall()]
        print(cart_items)

        user = c.execute('SELECT * FROM Users WHERE user_id = ?', (user_id,)).fetchone()
        print(user)
        user = {'username': user[1]}
        # Считаем общую сумму заказа
        total_price = sum(item['total'] for item in cart_items)
        conn.commit()
        conn.close()
        return render_template('cart.html', cart_items=cart_items, total_price=total_price, current_user=user)
    else:
        refresh_token = request.cookies.get('refresh_token')
        user_id = verify_jwt(refresh_token)
        if user_id:
            new_access_token = generate_jwt(user_id)
            new_refresh_token = generate_refresh_token(user_id)
            conn = get_connection('building_materials_shop.db')
            c = conn.cursor()
            c.execute('UPDATE Users SET refresh_token = ? WHERE user_id = ?', (new_refresh_token, user_id))
            conn.commit()
            conn.close()
            response = make_response(redirect(url_for('cart')))
            response.set_cookie('token', new_access_token, httponly=True)
            response.set_cookie('refresh_token', new_refresh_token, httponly=True)
            return response
        else:
            return render_template('login.html')


@app.route('/checkout', methods=['GET', 'POST'])
def checkout():
    global nums
    token = request.cookies.get('token')
    user_id = verify_jwt(token)
    if user_id:
        user_id = session['user_id']
        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()

        nums = c.execute('SELECT * FROM Cart WHERE user_id = ?', (user_id, )).fetchall()
        # Получаем все товары в корзине пользователя
        c.execute('SELECT p.type, p.brand, p.model, p.price, c.quantity, (p.price * c.quantity) as total, p.product_id '
                  'FROM Cart c '
                  'JOIN Products p ON c.product_id = p.product_id '
                  'WHERE c.user_id = ?', (user_id,))
        cart_items = [{'type': row[0], 'brand': row[1], 'model': row[2], 'price': row[3], 'quantity': row[4], 'total': row[5], 'id': row[6]} for row in c.fetchall()]

        # Считаем общую сумму заказа
        total_price = sum(item['total'] for item in cart_items)
        print(len(cart_items))
        if len(cart_items) == 0:
            return redirect(url_for('index'))

        if request.method == 'POST':
            name = request.form['name']
            address = request.form['address']
            phone = request.form['phone']
            payment_method = request.form['payment_method']

            # Сохраняем информацию о заказе в таблице Orders
           # c.execute('INSERT INTO Orders (user_id, name, address, phone, payment_method) VALUES (?, ?, ?, ?, ?)', (user_id, name, address, phone, payment_method))

            for item in cart_items:
                p_id = c.execute("SELECT * FROM Orders ORDER BY order_id DESC LIMIT 1").fetchone()
                position = p_id[0] +1
                c.execute('INSERT INTO Orders (user_id, name, address, phone, payment_method) VALUES (?, ?, ?, ?, ?)',
                          (user_id, name, address, phone, payment_method))
                c.execute('UPDATE Orders SET product_id = ?, quantity = ? WHERE order_id = ?',
                          (item['id'], item['quantity'], position))
            c.execute('DELETE FROM Cart WHERE user_id = ?', (user_id,))

            c.execute("UPDATE Products SET quantity = quantity - ? WHERE product_id = ?", (item['quantity'], item['id']))
            c.execute("UPDATE Products SET isExist = 0 WHERE quantity = 0")
            conn.commit()
            conn.close()


            if payment_method == 'cash':
                # Отображаем модальное окно
                return render_template('checkout.html', cart_items=cart_items, total_price=total_price,
                                       show_success_modal=True)
            else:
                return redirect(url_for('order_confirmation'))

        return render_template('checkout.html', cart_items=cart_items, total_price=total_price)
    else:
        refresh_token = request.cookies.get('refresh_token')
        user_id = verify_jwt(refresh_token)
        if user_id:
            new_access_token = generate_jwt(user_id)
            new_refresh_token = generate_refresh_token(user_id)
            conn = get_connection('building_materials_shop.db')
            c = conn.cursor()
            c.execute('UPDATE Users SET refresh_token = ? WHERE user_id = ?', (new_refresh_token, user_id))
            conn.commit()
            conn.close()
            response = make_response(redirect(url_for('checkout')))
            response.set_cookie('token', new_access_token, httponly=True)
            response.set_cookie('refresh_token', new_refresh_token, httponly=True)
            return response
        else:
            return render_template('login.html')


nums = []

@app.route('/order_confirmation')
def order_confirmation():
    if 'user_id' in session:
        user_id = session['user_id']
        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()

        print(len(nums))

        # Получаем последний заказ пользователя
        c.execute('SELECT p.type, p.brand, p.model, p.price, p.imageURL, o.quantity, (p.price * o.quantity) as total '
                  'FROM Orders o '
                  'JOIN Products p ON o.product_id = p.product_id '
                  'WHERE o.user_id = ? '
                  'ORDER BY o.order_id DESC LIMIT ?',(user_id, len(nums)))
        order_items = [
            {'type': row[0], 'brand': row[1], 'model': row[2], 'price': row[3], 'imageURL': row[4], 'quantity': row[5],
             'total': row[6]} for row in c.fetchall()]
        print(order_items)

        # Получаем информацию о доставке для последнего заказа
        c.execute('SELECT name, address, phone FROM Orders WHERE user_id = ? ORDER BY order_id DESC LIMIT 1', (user_id,))
        order_info = c.fetchone()



        conn.commit()
        conn.close()



        # Создаем документ Word с информацией о заказе
        document = Document()
        document.add_heading('Информация о заказе', 0)

        document.add_paragraph(f'Имя: {order_info[0]}')
        document.add_paragraph(f'Адрес: {order_info[1]}')
        document.add_paragraph(f'Телефон: {order_info[2]}')

        document.add_heading('Товары в заказе', level=1)
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Тип'
        hdr_cells[1].text = 'Бренд'
        hdr_cells[2].text = 'Модель'
        hdr_cells[3].text = 'Цена'
        hdr_cells[4].text = 'Количество'
        hdr_cells[5].text = 'Итоговая сумма'

        for item in order_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item['type']
            row_cells[1].text = item['brand']
            row_cells[2].text = item['model']
            row_cells[3].text = str(item['price'])
            row_cells[4].text = str(item['quantity'])
            row_cells[5].text = str(item['total'])

        # Сохраняем документ в байтовый поток
        byte_io = io.BytesIO()
        document.save(byte_io)
        byte_io.seek(0)

        # Отдаем файл Word пользователю
        return send_file(byte_io, download_name=f'order_{user_id}.docx', as_attachment=True)
    else:
        return redirect(url_for('login'))


@app.route('/confirmation')
def confirmation():
    token = request.cookies.get('token')
    user_id = verify_jwt(token)
    if user_id:

        try:
            # Перенаправление на главную страницу
            flash('Ваш заказ успешно оформлен!', 'success')
            return redirect(url_for('index'))

        except sqlite3.Error as e:
            flash('Произошла ошибка при оформлении заказа. Пожалуйста, попробуйте еще раз.', 'danger')
            return redirect(url_for('cart'))
    else:
        return redirect(url_for('login'))

@app.route('/delete_from_cart', methods=['POST'])
def delete_from_cart():
    user_id = request.form.get('user_id')
    product_id = request.form.get('product_id')
    print(f"User ID: {user_id}, Cart ID: {product_id}")
    conn = get_connection('building_materials_shop.db')
    c = conn.cursor()

    # Проверяем, что товар принадлежит данному пользователю
    c.execute("SELECT * FROM Cart WHERE user_id = ? AND product_id = ?", (user_id, product_id))
    if c.fetchone() is None:
        # Если товар не найден, возвращаем ошибку
        conn.close()
        return jsonify({"error": "Товар не найден в корзине"}), 404

    # Удаляем товар из корзины
    c.execute("DELETE FROM Cart WHERE user_id = ? AND product_id = ?", (user_id, product_id))
    conn.commit()
    conn.close()

    # Возвращаем успешный ответ
    return jsonify({"message": "Товар успешно удален из корзины"})


@app.route('/search')
def search():
    query = request.args.get('query')
    conn = get_connection('building_materials_shop.db')
    c = conn.cursor()
    c.execute("SELECT * FROM Products WHERE type LIKE ? OR brand LIKE ? OR model LIKE ?", (f"%{query}%", f"%{query}%", f"%{query}%"))
    products = [Product(row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7]) for row in c.fetchall()]
    conn.close()
    return render_template('index.html', products=products)


@app.route('/logout')
def logout():
    response = make_response(redirect(url_for('index')))
    response.set_cookie('token', '', expires=0)
    response.set_cookie('refresh_token', '', expires=0)
    return response

@app.route('/refresh', methods=['POST'])
def refresh():
    refresh_token = request.cookies.get('refresh_token')
    user_id = verify_jwt(refresh_token)
    if user_id:
        conn = get_connection('building_materials_shop.db')
        c = conn.cursor()
        user = c.execute('SELECT * FROM Users WHERE user_id = ?', (user_id,)).fetchone()
        if user and user[4] == refresh_token:  # Check if refresh token matches the one in DB
            new_access_token = generate_jwt(user_id)
            new_refresh_token = generate_refresh_token(user_id)
            c.execute('UPDATE Users SET refresh_token = ? WHERE user_id = ?', (new_refresh_token, user_id))
            conn.commit()
            conn.close()
            response = make_response(jsonify({'access_token': new_access_token, 'refresh_token': new_refresh_token}))
            response.set_cookie('token', new_access_token, httponly=True)
            response.set_cookie('refresh_token', new_refresh_token, httponly=True)
            return response
        else:
            return jsonify({'error': 'Invalid refresh token'}), 401
    else:
        return jsonify({'error': 'Refresh token expired or invalid'}), 401



@app.route('/change_password', methods=['POST'])
def change_password():
    access_token = request.cookies.get('token')
    user_id = verify_jwt(access_token)
    print(user_id)
    current_password = request.form['current_password']
    new_password = request.form['new_password']
    confirm_password = request.form['confirm_password']

    if new_password != confirm_password:
        flash('Новый пароль и подтверждение пароля не совпадают', 'danger')
        return redirect(url_for('profile'))

    conn = sqlite3.connect('building_materials_shop.db')
    c = conn.cursor()

    c.execute("SELECT * FROM users WHERE user_id = ?", (user_id, ))
    user = c.fetchone()
    if user and user[2] != current_password:
        flash("Вы ввели неправильный текущий пароль", "danger")
        conn.close()
        return redirect(url_for('profile'))

    c.execute("UPDATE users SET password = ? WHERE user_id = ?", (new_password, user_id))
    conn.commit()
    conn.close()
    flash("Пароль усешно изменён", "success")
    return redirect(url_for('profile'))



if __name__ == '__main__':
    app.run(debug=True)